from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from openai import OpenAI
from docx import Document
import logging
import os
import uuid

app = Flask(__name__)
CORS(app)
logging.basicConfig(level=logging.INFO)

# Set your OpenAI API key here or use environment variables
api_key = " "  # Replace with your actual key
client = OpenAI(api_key=api_key)
# Define your desired output folder (change this to your target path)
OUTPUT_DIR = "/Users/prudhvitalluri/Documents" # Replace with your location  
os.makedirs(OUTPUT_DIR, exist_ok=True)
@app.route('/generate_resume', methods=['POST'])
def generate_resume():
    data = request.get_json()
    name = data.get('name')
    title = data.get('title')
    in_message = data.get('in_message')

    if not all([name, title, in_message]):
        return jsonify({'error': 'Missing one of: name, title, or in_message'}), 400

    try:
        # Generate summary from OpenAI
        completion = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "You are an AI assistant that generates clean, professional, and ATS-optimized resumes. "
            "Use only the information provided as input. Do not include any placeholder fields such as "
            "[Your Name], [Address], [Email], [Phone Number], or [LinkedIn Profile].\n\n"
            "Resume formatting requirements:\n"
            "• All section headings (e.g., Professional Summary, Skills, Experience) must be bold.\n"
            "• Only include a section if relevant content is provided — no empty subheadings.\n"
            "• The Professional Summary must be a minimum of 10 lines, written in a professional and formal tone, "
            "highlighting domain expertise, technical skills, accomplishments, and leadership ability.\n"
            "• Each job experience entry under the Experience section must include at least 10 lines of detailed, "
            "results-driven bullet points or sentences describing responsibilities, achievements, technologies used, "
            "and business impact.\n"
            "• Resume should be well-structured, easy to read, and export-ready in PDF format.\n"
            "• Do not fabricate or assume any information that is not part of the input."
        },
                {"role": "user", "content": in_message}
            ]
        )
        summary = completion.choices[0].message.content.strip()

        # Create Word document
        doc = Document()
        doc.add_heading(name, 0)
        doc.add_heading(title, level=1)
        doc.add_paragraph('Summary:')
        doc.add_paragraph(summary)

       # Save file to specific directory
        file_id = uuid.uuid4().hex
        filename = f"resume_{file_id}.docx"
        full_path = os.path.join(OUTPUT_DIR, filename)
        doc.save(full_path)

        # Send file as download
        return send_file(full_path, as_attachment=True)


    except Exception as e:
        logging.error(f"Error: {str(e)}")
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    app.run(debug=True)